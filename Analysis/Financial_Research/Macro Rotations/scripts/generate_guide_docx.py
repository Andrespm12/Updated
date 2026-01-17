
import os
import re
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def markdown_to_docx(md_file, docx_file):
    document = Document()
    
    # Set default style to something clean
    style = document.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    with open(md_file, 'r') as f:
        lines = f.readlines()

    current_paragraph = None

    for line in lines:
        line = line.rstrip()
        
        # Headers
        if line.startswith('# '):
            document.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            document.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            document.add_heading(line[4:], level=3)
        elif line.startswith('#### '):
            document.add_heading(line[5:], level=4)
            
        # Horizontal Rule
        elif line.startswith('---') or line.startswith('***'):
            document.add_page_break()
            
        # Lists
        elif line.strip().startswith('- ') or line.strip().startswith('* '):
            # Indentation level
            indent = (len(line) - len(line.lstrip())) // 4
            content = line.strip()[2:]
            
            # Simple bold parsing **text**
            parts = re.split(r'(\*\*.*?\*\*)', content)
            
            p = document.add_paragraph(style='List Bullet')
            if indent > 0:
                p.paragraph_format.left_indent = Pt(18 * (indent + 1))
                
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    p.add_run(part)

        # Empty lines
        elif not line.strip():
            continue
            
        # Normal Text
        else:
            p = document.add_paragraph()
            parts = re.split(r'(\*\*.*?\*\*)', line)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    p.add_run(part)

    document.save(docx_file)
    print(f"Successfully converted {md_file} to {docx_file}")

if __name__ == "__main__":
    md_path = "/Users/andrespena/Desktop/Macro Rotations/Dashboard_Metrics_Guide.md"
    docx_path = "/Users/andrespena/Desktop/Macro Rotations/Dashboard_Metrics_Guide.docx"
    
    if os.path.exists(md_path):
        markdown_to_docx(md_path, docx_path)
    else:
        print(f"Error: {md_path} not found.")
